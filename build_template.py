"""
Genera la plantilla Word professional para informes de ventas.
Ejecutar una vez: python build_template.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def build():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    for lvl, sz in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f'Heading {lvl}']
        s.font.name = 'Calibri'
        s.font.size = Pt(sz)
        s.font.color.rgb = RGBColor(79, 70, 229)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('INFORME DE VENTAS')
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(79, 70, 229)
    r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('{{ asesor_name }}')
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(30, 41, 59)
    r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Canal: {{ canal }}')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('_' * 50)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Generado el {{ fecha }}')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    doc.add_heading('Resumen de Metricas', level=1)

    t = doc.add_table(rows=5, cols=4, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ('Cant. Pedida', '{{ cant_pedida }}', 'Cant. Pendiente', '{{ cant_pendiente }}'),
        ('Cant. Comprom.', '{{ cant_comprometida }}', 'Backlog', '{{ backlog }}'),
        ('Valor Total', '{{ valor_total }}', 'Utilidad Prom.', '{{ utilidad_promedio }}'),
        ('Margen Prom.', '{{ margen_promedio }}', 'Descuentos', '{{ descuentos }}'),
        ('Docs Unicos', '{{ docs_unicos }}', 'Total Registros', '{{ total_registros }}'),
    ]
    for ri, row_data in enumerate(kpis):
        for ci in range(4):
            cell = t.rows[ri].cells[ci]
            shade(cell, 'F8FAFC' if ri % 2 == 0 else 'FFFFFF')
            p = cell.paragraphs[0]
            r = p.add_run(row_data[ci])
            if ci % 2 == 0:
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(100, 116, 139)
                r.bold = True
            else:
                r.font.size = Pt(14)
                r.bold = True

    doc.add_page_break()

    doc.add_heading('Proyectos Principales', level=1)

    for section in ['chart_proyectos']:
        p = doc.add_paragraph()
        p.add_run('{{ ' + section + ' }}')

    doc.add_heading('Distribucion por Estado', level=1)

    for section in ['chart_estados']:
        p = doc.add_paragraph()
        p.add_run('{{ ' + section + ' }}')

    doc.add_heading('Tipo de Contrato', level=1)

    doc.add_page_break()

    doc.add_heading('Analisis con Inteligencia Artificial', level=1)

    p = doc.add_paragraph()
    p.add_run('{{ informe_content }}')

    doc.add_page_break()

    doc.add_heading('Graficos de Analisis', level=1)

    for section in ['chart_unidades']:
        p = doc.add_paragraph()
        p.add_run('{{ ' + section + ' }}')

    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'informe_template.docx')
    doc.save(out_path)
    print(f'Template saved: {out_path}')


if __name__ == '__main__':
    build()
