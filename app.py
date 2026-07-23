import os
import io
import json
import shutil
import tempfile
import numpy as np
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException, Header, Query, Body
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from analyzer import process_excel, build_asesor_prompt


def to_serializable(obj):
    if isinstance(obj, dict):
        return {k: to_serializable(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [to_serializable(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        return float(obj)
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    return obj


GEMINI_MODELS = ["gemini-3.1-flash-lite", "gemini-3-flash-preview", "gemini-3.5-flash"]
CLAUDE_MODELS = ["claude-sonnet-4-20250514", "claude-3-5-haiku-20241022"]
OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini"]


def _call_gemini(prompt: str, api_key: str) -> str:
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    last_error = None
    for model_name in GEMINI_MODELS:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            last_error = e
            err_str = str(e).lower()
            if "429" in err_str or "quota" in err_str or "rate" in err_str:
                continue
            raise e
    raise Exception(f"Todos los modelos Gemini fallaron. Ultimo error: {last_error}")


def _call_claude(prompt: str, api_key: str) -> str:
    import time
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    last_error = None
    for model_name in CLAUDE_MODELS:
        for attempt in range(3):
            try:
                response = client.messages.create(
                    model=model_name,
                    max_tokens=8192,
                    messages=[{"role": "user", "content": prompt}],
                )
                return response.content[0].text
            except Exception as e:
                last_error = e
                err_str = str(e).lower()
                if "429" in err_str or "rate" in err_str or "overloaded" in err_str:
                    if attempt < 2:
                        time.sleep(5 * (attempt + 1))
                        continue
                    else:
                        break
                else:
                    raise e
    raise Exception(f"Todos los modelos Claude fallaron. Ultimo error: {last_error}")


def _call_openai(prompt: str, api_key: str) -> str:
    import time
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    last_error = None
    for model_name in OPENAI_MODELS:
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=8192,
                )
                return response.choices[0].message.content
            except Exception as e:
                last_error = e
                err_str = str(e).lower()
                if "429" in err_str or "rate" in err_str or "quota" in err_str:
                    if attempt < 2:
                        time.sleep(5 * (attempt + 1))
                        continue
                    else:
                        break
                else:
                    raise e
    raise Exception(f"Todos los modelos OpenAI fallaron. Ultimo error: {last_error}")


def _call_ai(prompt: str, api_key: str, provider: str = "gemini") -> str:
    if provider == "claude":
        return _call_claude(prompt, api_key)
    elif provider == "openai":
        return _call_openai(prompt, api_key)
    else:
        return _call_gemini(prompt, api_key)


CHART_COLORS = ['#4F46E5', '#7C3AED', '#2563EB', '#3B82F6', '#8B5CF6',
                '#A78BFA', '#C4B5FD', '#DDD6FE', '#6366F1', '#818CF8']

def _chart_top_proyectos(asesor_data: dict) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    proyectos = asesor_data.get('top_proyectos', [])[:10]
    if not proyectos:
        return None

    nombres = [p['proyecto'][:35] for p in reversed(proyectos)]
    valores = [p['valor_pendiente'] for p in reversed(proyectos)]
    comprometido = [p['v_comprometido'] for p in reversed(proyectos)]

    fig, ax = plt.subplots(figsize=(9, max(3.5, len(nombres) * 0.55)))
    y_pos = range(len(nombres))

    bars1 = ax.barh(y_pos, valores, height=0.35, label='Valor Pendiente',
                     color=CHART_COLORS[0], alpha=0.85, edgecolor='white', linewidth=0.5)
    bars2 = ax.barh([y + 0.35 for y in y_pos], comprometido, height=0.35,
                     label='V. Comprometido', color=CHART_COLORS[2], alpha=0.85,
                     edgecolor='white', linewidth=0.5)

    ax.set_yticks([y + 0.175 for y in y_pos])
    ax.set_yticklabels(nombres, fontsize=8)
    ax.set_xlabel('Valor ($)', fontsize=9)
    ax.set_title('Top Proyectos por Valor', fontsize=12, fontweight='bold',
                  color='#1E293B', pad=12)
    ax.legend(fontsize=8, loc='lower right')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.2, linestyle='--')
    ax.tick_params(axis='both', labelsize=8)

    for bar in bars1:
        width = bar.get_width()
        if width > 0:
            ax.text(width + max(valores) * 0.01, bar.get_y() + bar.get_height()/2,
                     f'${width:,.0f}', va='center', fontsize=7, color='#475569')
    for bar in bars2:
        width = bar.get_width()
        if width > 0:
            ax.text(width + max(comprometido) * 0.01, bar.get_y() + bar.get_height()/2,
                     f'${width:,.0f}', va='center', fontsize=7, color='#475569')

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_estados_pedido(asesor_data: dict) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    estados = asesor_data.get('pedidos_por_estado', {})
    if not estados:
        return None

    labels = list(estados.keys())
    sizes = [estados[e]['registros'] for e in labels]
    colors_map = {
        'Aprobado': '#34D399', 'Retenido': '#F87171', 'En proceso': '#FBBF24',
        'Despachado': '#60A5FA', 'Facturado': '#A78BFA', 'Pendiente': '#FDE68A',
        'Anulado': '#94A3B8', 'Parcialmente Despachado': '#38BDF8'
    }
    colors = [colors_map.get(l, '#94A3B8') for l in labels]

    fig, ax = plt.subplots(figsize=(6, 4))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct='%1.0f%%', startangle=90,
        colors=colors, pctdistance=0.75, labeldistance=1.15,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)
    )
    for t in texts:
        t.set_fontsize(8)
    for t in autotexts:
        t.set_fontsize(7)
        t.set_fontweight('bold')

    ax.set_title('Distribucion por Estado de Pedido', fontsize=11,
                  fontweight='bold', color='#1E293B', pad=12)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_unidades_negocio(asesor_data: dict) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    unidades = asesor_data.get('unidades_negocio', {})
    if not unidades:
        return None

    items = sorted(unidades.items(), key=lambda x: x[1]['cant_pedida'], reverse=True)[:8]
    nombres = [u[0][:25] for u in items]
    pedidas = [u[1]['cant_pedida'] for u in items]
    valores = [u[1]['valor'] for u in items]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    bars1 = ax1.bar(range(len(nombres)), pedidas, color=CHART_COLORS[0], alpha=0.85,
                     edgecolor='white', linewidth=0.5)
    ax1.set_xticks(range(len(nombres)))
    ax1.set_xticklabels(nombres, rotation=35, ha='right', fontsize=7)
    ax1.set_ylabel('Cant. Pedida', fontsize=8)
    ax1.set_title('Por Cantidad Pedida', fontsize=10, fontweight='bold', color='#1E293B')
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.grid(axis='y', alpha=0.2, linestyle='--')
    ax1.tick_params(axis='both', labelsize=7)
    for bar in bars1:
        h = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2, h + max(pedidas)*0.01,
                  f'{h:,.0f}', ha='center', va='bottom', fontsize=6, color='#475569')

    bars2 = ax2.bar(range(len(nombres)), valores, color=CHART_COLORS[2], alpha=0.85,
                     edgecolor='white', linewidth=0.5)
    ax2.set_xticks(range(len(nombres)))
    ax2.set_xticklabels(nombres, rotation=35, ha='right', fontsize=7)
    ax2.set_ylabel('Valor ($)', fontsize=8)
    ax2.set_title('Por Valor', fontsize=10, fontweight='bold', color='#1E293B')
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.grid(axis='y', alpha=0.2, linestyle='--')
    ax2.tick_params(axis='both', labelsize=7)
    for bar in bars2:
        h = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2, h + max(valores)*0.01,
                  f'${h:,.0f}', ha='center', va='bottom', fontsize=6, color='#475569')

    plt.suptitle('Unidades de Negocio', fontsize=12, fontweight='bold',
                  color='#1E293B', y=1.02)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _verify_ai_key(api_key: str, provider: str = "gemini") -> dict:
    models_tried = []
    try:
        if provider == "claude":
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model=CLAUDE_MODELS[0],
                max_tokens=10,
                messages=[{"role": "user", "content": "Responde solo: OK"}],
            )
            return {"success": True, "model": CLAUDE_MODELS[0], "message": "Conexion exitosa con " + CLAUDE_MODELS[0]}
        elif provider == "openai":
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model=OPENAI_MODELS[0],
                messages=[{"role": "user", "content": "Responde solo: OK"}],
                max_tokens=10,
            )
            return {"success": True, "model": OPENAI_MODELS[0], "message": "Conexion exitosa con " + OPENAI_MODELS[0]}
        else:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            last_err = ""
            for model_name in GEMINI_MODELS:
                models_tried.append(model_name)
                try:
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content("Responde solo: OK")
                    return {"success": True, "model": model_name, "message": "Conexion exitosa con " + model_name}
                except Exception as e:
                    last_err = str(e)
                    continue
            raise Exception(last_err or "Todos los modelos fallaron")
    except Exception as e:
        err_str = str(e).lower()
        if "429" in err_str or "quota" in err_str or "rate" in err_str:
            tried = ", ".join(models_tried) if models_tried else provider
            raise HTTPException(status_code=429, detail=f"Cuota agotada. Modelos intentados: {tried}. Espera o usa otra API key.")
        if "invalid" in err_str or "unauthorized" in err_str or "401" in err_str:
            raise HTTPException(status_code=400, detail="API key invalida para " + provider)
        raise HTTPException(status_code=400, detail=f"Error verificando {provider}: {str(e)}")

load_dotenv()

app = FastAPI(title="Analisis de Pedidos", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

current_data = {}


@app.get("/", response_class=HTMLResponse)
async def serve_index():
    html_path = Path("index.html")
    if html_path.exists():
        return HTMLResponse(content=html_path.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>index.html no encontrado</h1>", status_code=404)


@app.post("/verify-key")
async def verify_key(
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    api_key = x_api_key or ""
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")
    return _verify_ai_key(api_key, provider)


@app.post("/upload")
async def upload_excel(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No se envio ningun archivo. Selecciona un archivo Excel.")

    if not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Formato no valido. Solo se permiten archivos Excel (.xlsx)")

    filepath = UPLOAD_DIR / file.filename
    try:
        with open(filepath, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except PermissionError:
        raise HTTPException(status_code=500, detail="No se pudo guardar el archivo. Verifica que no este abierto en Excel y que tengas permisos de escritura.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al guardar el archivo: {str(e)}")

    if filepath.stat().st_size == 0:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="El archivo esta vacio. Selecciona un archivo Excel valido.")

    try:
        result = process_excel(str(filepath))
        current_data["result"] = to_serializable(result)
        current_data["filename"] = file.filename
        current_data["filepath"] = str(filepath)
    except ValueError as e:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"Error al procesar el archivo: {str(e)}")

    return {
        "success": True,
        "filename": file.filename,
        "total_rows": result["total_raw_rows"],
        "total_unfiltered_rows": result.get("total_unfiltered_rows", result["total_raw_rows"]),
        "asesores": len(result["asesor_metrics"]),
        "canal_filter": result["canal_filter"],
        "detected_columns": result["detected_columns"],
    }


@app.post("/refilter")
async def refilter(canal: str = Query(...)):
    if not current_data.get("filepath"):
        raise HTTPException(status_code=400, detail="No hay archivo cargado. Sube un Excel primero.")
    filepath = current_data["filepath"]
    if not os.path.exists(filepath):
        raise HTTPException(status_code=400, detail="El archivo original ya no existe. Vuelve a subirlo.")
    try:
        result = process_excel(str(filepath), canal_filter_override=canal)
        current_data["result"] = to_serializable(result)
        return {
            "success": True,
            "canal_filter": result["canal_filter"],
            "canal_dist_values": result["canal_dist_values"],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al re-filtrar: {str(e)}")


@app.get("/data")
async def get_data():
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")
    return current_data["result"]


@app.post("/analyze/{asesor_name}")
async def analyze_asesor(
    asesor_name: str,
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
    canal: Optional[str] = Query(None),
):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")

    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    result = current_data["result"]
    asesor_metrics = result["asesor_metrics"]
    team_summary = result["team_summary"]

    asesor_data = None
    for m in asesor_metrics:
        if m["asesor"].strip().upper() == asesor_name.strip().upper():
            asesor_data = m
            break

    if not asesor_data:
        raise HTTPException(status_code=404, detail=f"Asesor '{asesor_name}' no encontrado.")

    canal_filter = canal or result.get("canal_filter", "")
    prompt = build_asesor_prompt(asesor_data, team_summary, canal_filter)

    try:
        informe = _call_ai(prompt, api_key, provider)
    except Exception as e:
        err_str = str(e).lower()
        if "429" in err_str or "quota" in err_str or "rate" in err_str:
            raise HTTPException(status_code=429, detail="Cuota agotada. Verifica tu plan o usa otra API key.")
        raise HTTPException(status_code=500, detail=f"Error con {provider}: {str(e)}")

    cache_key = f"report_{asesor_name.strip().upper()}"
    current_data[cache_key] = informe

    return {
        "asesor": asesor_name,
        "informe": informe,
        "metricas": asesor_data,
    }


@app.post("/analyze-all")
async def analyze_all(
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
    canal: Optional[str] = Query(None),
):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")

    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    result = current_data["result"]
    asesor_metrics = result["asesor_metrics"]
    team_summary = result["team_summary"]
    canal_filter = canal or result.get("canal_filter", "")

    informes = []
    for m in asesor_metrics:
        try:
            prompt = build_asesor_prompt(m, team_summary, canal_filter)
            informe_text = _call_ai(prompt, api_key, provider)
            informes.append({
                "asesor": m["asesor"],
                "informe": informe_text,
                "metricas": m,
            })
        except Exception as e:
            informes.append({
                "asesor": m["asesor"],
                "informe": f"Error con {provider}: {str(e)}",
                "metricas": m,
            })

    return {"informes": informes}


@app.get("/export/{asesor_name}")
async def export_asesor(asesor_name: str):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")

    result = current_data["result"]
    asesor_data = None
    for m in result["asesor_metrics"]:
        if m["asesor"].strip().upper() == asesor_name.strip().upper():
            asesor_data = m
            break

    if not asesor_data:
        raise HTTPException(status_code=404, detail=f"Asesor '{asesor_name}' no encontrado.")

    return asesor_data


def _get_asesor_data(asesor_name: str) -> dict:
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")
    result = current_data["result"]
    for m in result["asesor_metrics"]:
        if m["asesor"].strip().upper() == asesor_name.strip().upper():
            return m
    raise HTTPException(status_code=404, detail=f"Asesor '{asesor_name}' no encontrado.")


def _generate_asesor_report(asesor_name: str, api_key: str, canal: str = "", provider: str = "gemini") -> str:
    result = current_data["result"]
    asesor_data = _get_asesor_data(asesor_name)
    team_summary = result["team_summary"]
    canal_filter = canal or result.get("canal_filter", "")

    cache_key = f"report_{asesor_name.strip().upper()}"
    if cache_key in current_data:
        return current_data[cache_key]

    prompt = build_asesor_prompt(asesor_data, team_summary, canal_filter)
    informe = _call_ai(prompt, api_key, provider)
    current_data[cache_key] = informe
    return informe


@app.post("/generate-and-cache/{asesor_name}")
async def generate_and_cache(
    asesor_name: str,
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
    canal: Optional[str] = Query(None),
):
    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    try:
        informe = _generate_asesor_report(asesor_name, api_key, canal, provider)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar informe: {str(e)}")

    return {"asesor": asesor_name, "informe": informe}


@app.get("/download/{asesor_name}/word")
async def download_word(
    asesor_name: str,
    canal: Optional[str] = Query(None),
):
    try:
        if not current_data or "result" not in current_data:
            return JSONResponse(status_code=400, content={"detail": "No hay datos cargados. Sube un archivo Excel primero."})
        asesor_data = _get_asesor_data(asesor_name)
        result = current_data["result"]
        canal_filter = canal or result.get("canal_filter", "")
        cache_key = f"report_{asesor_name.strip().upper()}"
        informe = current_data.get(cache_key, "")
    except Exception as e:
        return JSONResponse(status_code=400, content={"detail": f"Error al obtener datos: {str(e)}"})

    try:
        doc = _build_word_doc(asesor_name, canal_filter, informe, asesor_data)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_name = asesor_name.replace(' ', '_').replace('.', '').replace('/', '_')
        return StreamingResponse(
            buffer,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="Informe_{safe_name}.docx"'}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": f"Error al generar Word: {str(e)}"})


@app.get("/download/{asesor_name}/excel")
async def download_excel(
    asesor_name: str,
    canal: Optional[str] = Query(None),
):
    try:
        if not current_data or "result" not in current_data:
            return JSONResponse(status_code=400, content={"detail": "No hay datos cargados. Sube un archivo Excel primero."})
        asesor_data = _get_asesor_data(asesor_name)
        result = current_data["result"]
        canal_filter = canal or result.get("canal_filter", "")
    except Exception as e:
        return JSONResponse(status_code=400, content={"detail": f"Error al obtener datos: {str(e)}"})

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()

        header_font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
        header_fill = PatternFill(start_color='4F46E5', end_color='4F46E5', fill_type='solid')
        title_font = Font(name='Calibri', bold=True, size=16, color='4F46E5')
        subtitle_font = Font(name='Calibri', bold=True, size=10, color='64748B')
        metric_label_font = Font(name='Calibri', bold=True, size=11)
        metric_value_font = Font(name='Calibri', size=11)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        ws = wb.active
        ws.title = 'Metricas Asesor'

        ws.merge_cells('A1:D1')
        ws['A1'] = f'Informe de Ventas - {asesor_name}'
        ws['A1'].font = title_font
        ws['A1'].alignment = Alignment(horizontal='center')

        ws.merge_cells('A2:D2')
        ws['A2'] = f'Canal: {canal_filter}'
        ws['A2'].font = subtitle_font
        ws['A2'].alignment = Alignment(horizontal='center')

        ws['A4'] = 'Metrica'
        ws['B4'] = 'Valor'
        ws['A4'].font = header_font
        ws['A4'].fill = header_fill
        ws['B4'].font = header_font
        ws['B4'].fill = header_fill
        ws['A4'].border = thin_border
        ws['B4'].border = thin_border

        metrics_data = [
            ('Cantidad Pedida', asesor_data['cant_pedida']),
            ('Cantidad Pendiente', asesor_data['cant_pendiente']),
            ('Cantidad Comprometida', asesor_data['cant_comprometida']),
            ('Backlog (%)', asesor_data['backlog_pct']),
            ('Valor Total (V.UNIDAD)', asesor_data.get('valor_total', 0)),
            ('Utilidad Promedio', asesor_data.get('utilidad_promedio', 0)),
            ('Margen Promedio (%)', asesor_data.get('margen_promedio', 0)),
            ('Descuentos Totales', asesor_data.get('descuento_total', 0)),
            ('Documentos Unicos', asesor_data['documentos_unicos']),
            ('Pedidos Retenidos', asesor_data['pedidos_retenidos']),
            ('Pedidos Aprobados', asesor_data['pedidos_aprobados']),
            ('Total Registros', asesor_data['total_registros']),
        ]

        for i, (label, value) in enumerate(metrics_data):
            row = i + 5
            ws.cell(row=row, column=1, value=label).font = metric_label_font
            cell = ws.cell(row=row, column=2, value=value)
            cell.font = metric_value_font
            cell.number_format = '#,##0'
            ws.cell(row=row, column=1).border = thin_border
            ws.cell(row=row, column=2).border = thin_border

        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20

        if asesor_data.get('pedidos_por_estado'):
            ws2 = wb.create_sheet('Pedidos por Estado')
            headers2 = ['Estado', 'Registros', 'Cant. Pedida', 'Cant. Pendiente']
            for j, h in enumerate(headers2):
                cell = ws2.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, (estado, data) in enumerate(asesor_data['pedidos_por_estado'].items()):
                ws2.cell(row=i + 2, column=1, value=estado)
                ws2.cell(row=i + 2, column=2, value=data['registros'])
                ws2.cell(row=i + 2, column=3, value=data['cant_pedida'])
                ws2.cell(row=i + 2, column=4, value=data['cant_pendiente'])
            ws2.column_dimensions['A'].width = 30
            ws2.column_dimensions['B'].width = 12
            ws2.column_dimensions['C'].width = 15
            ws2.column_dimensions['D'].width = 15

        if asesor_data.get('top_lineas'):
            ws3 = wb.create_sheet('Top Lineas')
            headers = ['Linea', 'Cant. Pedida', 'Registros']
            for j, h in enumerate(headers):
                cell = ws3.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, linea in enumerate(asesor_data['top_lineas']):
                ws3.cell(row=i + 2, column=1, value=linea['linea'])
                ws3.cell(row=i + 2, column=2, value=linea['cant_pedida'])
                ws3.cell(row=i + 2, column=3, value=linea['registros'])
            ws3.column_dimensions['A'].width = 40
            ws3.column_dimensions['B'].width = 15
            ws3.column_dimensions['C'].width = 12

        if asesor_data.get('top_sub_lineas'):
            ws4 = wb.create_sheet('Top Sub-Lineas')
            headers = ['Sub-Linea', 'Cant. Pedida', 'Registros']
            for j, h in enumerate(headers):
                cell = ws4.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, sub in enumerate(asesor_data['top_sub_lineas']):
                ws4.cell(row=i + 2, column=1, value=sub['sub_linea'])
                ws4.cell(row=i + 2, column=2, value=sub['cant_pedida'])
                ws4.cell(row=i + 2, column=3, value=sub['registros'])
            ws4.column_dimensions['A'].width = 40
            ws4.column_dimensions['B'].width = 15
            ws4.column_dimensions['C'].width = 12

        if asesor_data.get('top_clientes'):
            ws5 = wb.create_sheet('Top Clientes')
            headers = ['Cliente', 'Cant. Pedida', 'Registros']
            for j, h in enumerate(headers):
                cell = ws5.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, cliente in enumerate(asesor_data['top_clientes']):
                ws5.cell(row=i + 2, column=1, value=cliente['cliente'])
                ws5.cell(row=i + 2, column=2, value=cliente['cant_pedida'])
                ws5.cell(row=i + 2, column=3, value=cliente['registros'])
            ws5.column_dimensions['A'].width = 50
            ws5.column_dimensions['B'].width = 15
            ws5.column_dimensions['C'].width = 12

        if asesor_data.get('unidades_negocio'):
            ws6 = wb.create_sheet('Unidades Negocio')
            headers = ['Unidad', 'Cant. Pedida', 'Valor', 'Registros']
            for j, h in enumerate(headers):
                cell = ws6.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, (co_name, co_data) in enumerate(asesor_data['unidades_negocio'].items()):
                ws6.cell(row=i + 2, column=1, value=co_name)
                ws6.cell(row=i + 2, column=2, value=co_data['cant_pedida'])
                ws6.cell(row=i + 2, column=3, value=co_data['valor'])
                ws6.cell(row=i + 2, column=4, value=co_data['registros'])
            ws6.column_dimensions['A'].width = 40
            ws6.column_dimensions['B'].width = 15
            ws6.column_dimensions['C'].width = 20
            ws6.column_dimensions['D'].width = 12

        if asesor_data.get('top_estados_produccion'):
            ws7 = wb.create_sheet('Estados Produccion')
            ws7['A1'] = 'Estado'
            ws7['B1'] = 'Cantidad'
            ws7['A1'].font = header_font
            ws7['A1'].fill = header_fill
            ws7['B1'].font = header_font
            ws7['B1'].fill = header_fill
            for i, (estado, cant) in enumerate(asesor_data['top_estados_produccion'].items()):
                ws7.cell(row=i + 2, column=1, value=estado)
                ws7.cell(row=i + 2, column=2, value=cant)
            ws7.column_dimensions['A'].width = 30
            ws7.column_dimensions['B'].width = 15

        if asesor_data.get('top_proyectos'):
            ws8 = wb.create_sheet('Proyectos')
            headers8 = ['Proyecto', 'Pedida', 'Pendiente', 'Comprometida', 'Valor Pend.', 'V.Comprometido']
            for j, h in enumerate(headers8):
                cell = ws8.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, p in enumerate(asesor_data['top_proyectos']):
                ws8.cell(row=i + 2, column=1, value=p['proyecto'])
                ws8.cell(row=i + 2, column=2, value=p['cant_pedida'])
                ws8.cell(row=i + 2, column=3, value=p['cant_pendiente'])
                ws8.cell(row=i + 2, column=4, value=p['cant_comprometida'])
                ws8.cell(row=i + 2, column=5, value=p['valor_pendiente'])
                ws8.cell(row=i + 2, column=6, value=p['v_comprometido'])
            ws8.column_dimensions['A'].width = 55
            ws8.column_dimensions['B'].width = 15
            ws8.column_dimensions['C'].width = 15
            ws8.column_dimensions['D'].width = 15
            ws8.column_dimensions['E'].width = 20
            ws8.column_dimensions['F'].width = 20

            docs_por_proy = asesor_data.get('documentos_por_proyecto', {})
            all_docs = []
            all_items = []
            for p in asesor_data['top_proyectos']:
                docs = docs_por_proy.get(p['proyecto'], [])
                for d in docs:
                    all_docs.append({**d, 'proyecto': p['proyecto']})
                    for it in d.get('items', []):
                        all_items.append({**it, 'documento': d['documento'], 'proyecto': p['proyecto']})

            if all_docs:
                ws_doc = wb.create_sheet('Documentos')
                doc_headers = ['Proyecto', 'Documento', 'Pedida', 'Pendiente', 'Comprometida', 'Valor Pend.', 'V.Comprometido']
                for j, h in enumerate(doc_headers):
                    cell = ws_doc.cell(row=1, column=j + 1, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                for i, d in enumerate(all_docs):
                    ws_doc.cell(row=i + 2, column=1, value=d['proyecto'])
                    ws_doc.cell(row=i + 2, column=2, value=d['documento'])
                    ws_doc.cell(row=i + 2, column=3, value=d['cant_pedida'])
                    ws_doc.cell(row=i + 2, column=4, value=d['cant_pendiente'])
                    ws_doc.cell(row=i + 2, column=5, value=d['cant_comprometida'])
                    ws_doc.cell(row=i + 2, column=6, value=d['valor_pendiente'])
                    ws_doc.cell(row=i + 2, column=7, value=d['v_comprometido'])
                ws_doc.column_dimensions['A'].width = 45
                ws_doc.column_dimensions['B'].width = 20
                ws_doc.column_dimensions['C'].width = 12
                ws_doc.column_dimensions['D'].width = 12
                ws_doc.column_dimensions['E'].width = 12
                ws_doc.column_dimensions['F'].width = 18
                ws_doc.column_dimensions['G'].width = 18

            if all_items:
                ws_it = wb.create_sheet('Items')
                item_headers = ['Proyecto', 'Documento', 'Item', 'Pedida', 'Pendiente', 'Comprometida', 'Valor Pend.', 'V.Comprometido']
                for j, h in enumerate(item_headers):
                    cell = ws_it.cell(row=1, column=j + 1, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                for i, it in enumerate(all_items):
                    ws_it.cell(row=i + 2, column=1, value=it['proyecto'])
                    ws_it.cell(row=i + 2, column=2, value=it['documento'])
                    ws_it.cell(row=i + 2, column=3, value=it['item'])
                    ws_it.cell(row=i + 2, column=4, value=it['cant_pedida'])
                    ws_it.cell(row=i + 2, column=5, value=it['cant_pendiente'])
                    ws_it.cell(row=i + 2, column=6, value=it['cant_comprometida'])
                    ws_it.cell(row=i + 2, column=7, value=it['valor_pendiente'])
                    ws_it.cell(row=i + 2, column=8, value=it['v_comprometido'])
                ws_it.column_dimensions['A'].width = 40
                ws_it.column_dimensions['B'].width = 18
                ws_it.column_dimensions['C'].width = 45
                ws_it.column_dimensions['D'].width = 12
                ws_it.column_dimensions['E'].width = 12
                ws_it.column_dimensions['F'].width = 12
                ws_it.column_dimensions['G'].width = 18
                ws_it.column_dimensions['H'].width = 18

        if asesor_data.get('desglose_contrato'):
            ws10 = wb.create_sheet('Tipo Contrato')
            headers10 = ['Tipo', 'Pedida', 'Pendiente', 'Comprometida', 'Registros']
            for j, h in enumerate(headers10):
                cell = ws10.cell(row=1, column=j + 1, value=h)
                cell.font = header_font
                cell.fill = header_fill
            for i, (tipo, d) in enumerate(asesor_data['desglose_contrato'].items()):
                ws10.cell(row=i + 2, column=1, value=tipo)
                ws10.cell(row=i + 2, column=2, value=d['pedida'])
                ws10.cell(row=i + 2, column=3, value=d['pendiente'])
                ws10.cell(row=i + 2, column=4, value=d['comprometida'])
                ws10.cell(row=i + 2, column=5, value=d['registros'])
            ws10.column_dimensions['A'].width = 20
            ws10.column_dimensions['B'].width = 15
            ws10.column_dimensions['C'].width = 15
            ws10.column_dimensions['D'].width = 15
            ws10.column_dimensions['E'].width = 12

        cache_key = f"report_{asesor_name.strip().upper()}"
        informe = current_data.get(cache_key, "")
        if informe:
            ws_report = wb.create_sheet('Informe de IA')
            lines = informe.split('\n')
            row = 1
            i = 0
            while i < len(lines):
                stripped = lines[i].strip()

                if stripped.startswith('# '):
                    ws_report.cell(row=row, column=1, value=stripped[2:]).font = Font(name='Calibri', bold=True, size=16, color='4F46E5')
                    ws_report.merge_cells(f'A{row}:F{row}')
                    row += 1
                elif stripped.startswith('## '):
                    ws_report.cell(row=row, column=1, value=stripped[3:]).font = Font(name='Calibri', bold=True, size=13, color='4F46E5')
                    ws_report.merge_cells(f'A{row}:F{row}')
                    row += 1
                elif stripped.startswith('### '):
                    ws_report.cell(row=row, column=1, value=stripped[4:]).font = Font(name='Calibri', bold=True, size=11, color='334155')
                    ws_report.merge_cells(f'A{row}:F{row}')
                    row += 1

                elif stripped.startswith('|') and '|' in stripped[1:]:
                    table_rows = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        row_line = lines[i].strip()
                        cells = [c.strip() for c in row_line.split('|')[1:-1]]
                        is_sep = all(c.replace('-', '').replace(':', '').strip() == '' for c in cells)
                        if not is_sep:
                            table_rows.append(cells)
                        i += 1
                    i -= 1

                    if table_rows:
                        num_cols = max(len(r) for r in table_rows)
                        for ri, row_data in enumerate(table_rows):
                            for ci, cell_text in enumerate(row_data):
                                if ci < num_cols:
                                    cell = ws_report.cell(row=row, column=ci + 1, value=cell_text)
                                    if ri == 0:
                                        cell.font = header_font
                                        cell.fill = header_fill
                                    else:
                                        cell.font = Font(name='Calibri', size=10)
                                    cell.border = thin_border
                                    cell.alignment = Alignment(wrap_text=True)
                            row += 1
                        row += 1

                elif stripped.startswith('- '):
                    ws_report.cell(row=row, column=1, value='  •  ' + stripped[2:]).font = Font(name='Calibri', size=10)
                    row += 1
                elif stripped:
                    ws_report.cell(row=row, column=1, value=stripped).font = Font(name='Calibri', size=10)
                    row += 1
                i += 1

            ws_report.column_dimensions['A'].width = 30
            ws_report.column_dimensions['B'].width = 20
            ws_report.column_dimensions['C'].width = 20
            ws_report.column_dimensions['D'].width = 20
            ws_report.column_dimensions['E'].width = 20
            ws_report.column_dimensions['F'].width = 20

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        safe_name = asesor_name.replace(' ', '_').replace('.', '').replace('/', '_')
        return StreamingResponse(
            buffer,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="Informe_{safe_name}.xlsx"'}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": f"Error al generar Excel: {str(e)}"})


ADVISOR_EMAILS = {
    "INES SANCHEZ": "ines.sanchez@interdoors.com.co",
    "ELIANA": "eliana.gonzalez@interdoors.com.co",
    "KAROLIN": "karolin.gonzalez@interdoors.com.co",
    "MATEO POSADA": "mateo.posada@interdoors.com.co",
    "LEONARDO": "asesor3@interdoors.com.co",
    "YUDY CARRASQUILLA": "yudy.carrasquilla@interdoors.com.co",
    "LAURA OCHOA": "laura.ochoa@interdoors.com.co",
}


def _find_advisor_email(asesor_name: str) -> str:
    name_upper = asesor_name.strip().upper()
    if name_upper in ADVISOR_EMAILS:
        return ADVISOR_EMAILS[name_upper]
    for key, email in ADVISOR_EMAILS.items():
        if key in name_upper or name_upper in key:
            return email
        key_parts = key.split()
        if len(key_parts) >= 2 and all(p in name_upper for p in key_parts):
            return email
    return ""


def _send_email_smtp(to_email: str, subject: str, body_text: str, attachment_bytes: bytes = None, attachment_name: str = None):
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    host = os.getenv("SMTP_HOST", "")
    port = int(os.getenv("SMTP_PORT", "587"))
    user = os.getenv("SMTP_USER", "")
    password = os.getenv("SMTP_PASS", "")
    from_addr = os.getenv("SMTP_FROM", user)

    if not all([host, user, password]):
        raise Exception("Variables SMTP_HOST, SMTP_USER, SMTP_PASS no configuradas.")

    msg = MIMEMultipart()
    msg["From"] = from_addr
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    if attachment_bytes and attachment_name:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment_bytes)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{attachment_name}"')
        msg.attach(part)

    with smtplib.SMTP(host, port, timeout=30) as server:
        server.starttls()
        server.login(user, password)
        server.sendmail(from_addr, [to_email], msg.as_string())


def _test_smtp_connection():
    import smtplib
    host = os.getenv("SMTP_HOST", "")
    port = int(os.getenv("SMTP_PORT", "587"))
    user = os.getenv("SMTP_USER", "")
    password = os.getenv("SMTP_PASS", "")
    if not all([host, user, password]):
        raise Exception("Variables SMTP no configuradas.")
    with smtplib.SMTP(host, port, timeout=15) as server:
        server.starttls()
        server.login(user, password)
    return user


def _send_email(to_email: str, subject: str, body_text: str, attachment_bytes: bytes = None, attachment_name: str = None):
    import base64
    import urllib.request
    import json as _json

    resend_key = os.getenv("RESEND_API_KEY", "")

    if not resend_key and os.getenv("SMTP_HOST"):
        _send_email_smtp(to_email, subject, body_text, attachment_bytes, attachment_name)
        return

    if not resend_key:
        raise Exception("No hay RESEND_API_KEY ni SMTP configurado.")

    payload = {
        "from": "Analisis de Pedidos <practicante.comercial@interdoors.com.co>",
        "to": [to_email],
        "subject": subject,
        "text": body_text,
    }

    if attachment_bytes and attachment_name:
        b64 = base64.b64encode(attachment_bytes).decode("utf-8")
        payload["attachments"] = [{
            "filename": attachment_name,
            "content": b64,
        }]

    data = _json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        "https://api.resend.com/emails",
        data=data,
        headers={
            "Authorization": f"Bearer {resend_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = _json.loads(resp.read().decode("utf-8"))
            if resp.status not in (200, 201):
                raise Exception(f"Resend error: {result}")
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8") if e.fp else ""
        raise Exception(f"Resend HTTP {e.code}: {body}")


ACCENT_COLOR = RGBColor(79, 70, 229)
HEADER_BG = '4F46E5'
HEADER_FG = RGBColor(255, 255, 255)
ALT_ROW_BG = 'F1F5F9'
SUBTITLE_COLOR = RGBColor(100, 116, 139)
DARK_TEXT = RGBColor(30, 41, 59)
GREEN_COLOR = RGBColor(34, 197, 94)
RED_COLOR = RGBColor(239, 68, 68)
YELLOW_COLOR = RGBColor(234, 179, 8)


def _style_header_row(row, num_cols, font_size=9):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for ci in range(num_cols):
        cell = row.cells[ci]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), HEADER_BG)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                run.font.color.rgb = HEADER_FG
                run.bold = True


def _style_alt_rows(table, start_row=1, font_size=9):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for ri in range(start_row, len(table.rows)):
        if ri % 2 == 0:
            for cell in table.rows[ri].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), ALT_ROW_BG)
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
        for cell in table.rows[ri].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = DARK_TEXT


def _add_bold_text(paragraph, text, font_size=11, color=None):
    parts = text.split('**')
    for j, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        run.bold = (j % 2 == 1)
        if color:
            run.font.color.rgb = color


def _add_markdown_to_doc(doc, informe):
    if not informe:
        return
    lines = informe.split('\n')
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            _add_bold_text(p, stripped[5:], font_size=11, color=DARK_TEXT)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)

        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)

        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)

        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)

        elif stripped.startswith('|') and '|' in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                is_sep = all(c.replace('-', '').replace(':', '').strip() == '' for c in cells)
                if not is_sep:
                    table_rows.append(cells)
                i += 1
            i -= 1

            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=num_cols, style='Table Grid')
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            t.rows[ri].cells[ci].text = cell_text
                _style_header_row(t.rows[0], num_cols, font_size=9)
                _style_alt_rows(t, start_row=1, font_size=9)

        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_text(p, stripped[2:], font_size=10)

        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.)':
            p = doc.add_paragraph(style='List Number')
            _add_bold_text(p, stripped[2:].strip() if stripped[1] == '.' else stripped[3:].strip(), font_size=10)

        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped[2:])
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = SUBTITLE_COLOR

        elif stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(203, 213, 225)

        elif stripped:
            p = doc.add_paragraph()
            _add_bold_text(p, stripped, font_size=10)

        i += 1


def _add_title_page(doc, asesor_name, canal_filter):
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('INFORME DE VENTAS')
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT_COLOR
    run.bold = True

    doc.add_paragraph()

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(asesor_name)
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_TEXT
    run.bold = True

    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run(f'Canal: {canal_filter}')
    run.font.size = Pt(12)
    run.font.color.rgb = SUBTITLE_COLOR

    doc.add_paragraph()

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('─' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    run = footer_p.add_run(f'Generado el {datetime.now().strftime("%d/%m/%Y")}')
    run.font.size = Pt(10)
    run.font.color.rgb = SUBTITLE_COLOR


def _add_kpi_summary(doc, asesor_data):
    doc.add_heading('Resumen de Metricas', level=1)
    table = doc.add_table(rows=5, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    kpis = [
        ('Cant. Pedida', f"{asesor_data['cant_pedida']:,.0f}", GREEN_COLOR),
        ('Cant. Pendiente', f"{asesor_data['cant_pendiente']:,.0f}", YELLOW_COLOR),
        ('Cant. Comprometida', f"{asesor_data['cant_comprometida']:,.0f}", ACCENT_COLOR),
        ('Backlog', f"{asesor_data['backlog_pct']}%", RED_COLOR if asesor_data['backlog_pct'] > 70 else GREEN_COLOR),
        ('Valor Total', f"${asesor_data.get('valor_total', 0):,.0f}", ACCENT_COLOR),
        ('Utilidad Prom.', f"${asesor_data.get('utilidad_promedio', 0):,.0f}", GREEN_COLOR),
        ('Margen Prom.', f"{asesor_data.get('margen_promedio', 0):.0f}%", GREEN_COLOR),
        ('Descuentos', f"${asesor_data.get('descuento_total', 0):,.0f}", RED_COLOR if asesor_data.get('descuento_total', 0) > 0 else GREEN_COLOR),
        ('Docs Unicos', str(asesor_data['documentos_unicos']), ACCENT_COLOR),
        ('Total Registros', str(asesor_data['total_registros']), ACCENT_COLOR),
        ('', '', None),
        ('', '', None),
    ]

    for ri in range(5):
        for ci in range(2):
            kpi_idx = ri * 2 + ci
            if kpi_idx < len(kpis):
                label, value, color = kpis[kpi_idx]
                if label:
                    label_cell = table.rows[ri].cells[ci * 2]
                    value_cell = table.rows[ri].cells[ci * 2 + 1]
                    p_label = label_cell.paragraphs[0]
                    r_label = p_label.add_run(label)
                    r_label.font.size = Pt(8)
                    r_label.font.color.rgb = SUBTITLE_COLOR
                    r_label.bold = True
                    p_value = value_cell.paragraphs[0]
                    r_value = p_value.add_run(value)
                    r_value.font.size = Pt(14)
                    r_value.bold = True
                    if color:
                        r_value.font.color.rgb = color

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for ri in range(5):
        for ci in range(4):
            cell = table.rows[ri].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F8FAFC' if ri % 2 == 0 else 'FFFFFF')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)


def _build_word_doc(asesor_name, canal_filter, informe, asesor_data):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from datetime import datetime

    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates', 'informe_template.docx')

    use_template = os.path.exists(template_path)
    if use_template:
        try:
            from docxtpl import DocxTemplate
            tpl = DocxTemplate(template_path)
            context = {
                'asesor_name': asesor_name,
                'canal': canal_filter,
                'fecha': datetime.now().strftime("%d/%m/%Y"),
                'cant_pedida': f"{asesor_data['cant_pedida']:,.0f}",
                'cant_pendiente': f"{asesor_data['cant_pendiente']:,.0f}",
                'cant_comprometida': f"{asesor_data['cant_comprometida']:,.0f}",
                'backlog': f"{asesor_data['backlog_pct']}%",
                'valor_total': f"${asesor_data.get('valor_total', 0):,.0f}",
                'utilidad_promedio': f"${asesor_data.get('utilidad_promedio', 0):,.0f}",
                'margen_promedio': f"{asesor_data.get('margen_promedio', 0):.0f}%",
                'descuentos': f"${asesor_data.get('descuento_total', 0):,.0f}",
                'docs_unicos': str(asesor_data['documentos_unicos']),
                'total_registros': str(asesor_data['total_registros']),
                'chart_proyectos': '',
                'chart_estados': '',
                'chart_unidades': '',
                'informe_content': '',
            }
            buf = io.BytesIO()
            tpl.render(context)
            tpl.save(buf)
            buf.seek(0)
            doc = Document(buf)
        except Exception:
            use_template = False

    if not use_template:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        _add_title_page(doc, asesor_name, canal_filter)

    _add_kpi_summary(doc, asesor_data)

    sections_available = []
    if asesor_data.get('pedidos_por_estado'):
        sections_available.append('Distribucion por Estado')
    if asesor_data.get('top_lineas'):
        sections_available.append('Lineas de Producto')
    if asesor_data.get('unidades_negocio'):
        sections_available.append('Unidades de Negocio')
    if asesor_data.get('top_estados_produccion'):
        sections_available.append('Estados de Produccion')
    if asesor_data.get('top_proyectos'):
        sections_available.append('Proyectos')
    if asesor_data.get('desglose_contrato'):
        sections_available.append('Tipo de Contrato')
    if informe:
        sections_available.append('Analisis con Inteligencia Artificial')

    if sections_available:
        doc.add_heading('Indice', level=1)
        for idx, sec_name in enumerate(sections_available, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{idx}.  {sec_name}')
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_TEXT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        doc.add_page_break()

    if asesor_data.get('pedidos_por_estado'):
        doc.add_heading('Distribucion por Estado', level=1)
        estado_table = doc.add_table(rows=len(asesor_data['pedidos_por_estado']) + 1, cols=4, style='Table Grid')
        estado_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, h in enumerate(['Estado', 'Registros', 'Cant. Pedida', 'Cant. Pendiente']):
            estado_table.rows[0].cells[idx].text = h
        _style_header_row(estado_table.rows[0], 4, font_size=9)
        for i, (estado, data) in enumerate(asesor_data['pedidos_por_estado'].items()):
            row = estado_table.rows[i + 1]
            row.cells[0].text = estado
            row.cells[1].text = str(data['registros'])
            row.cells[2].text = f"{data['cant_pedida']:,.0f}"
            row.cells[3].text = f"{data['cant_pendiente']:,.0f}"
        _style_alt_rows(estado_table, start_row=1, font_size=9)

        chart_buf = _chart_estados_pedido(asesor_data)
        if chart_buf:
            doc.add_paragraph()
            doc.add_picture(chart_buf, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if asesor_data.get('top_lineas'):
        doc.add_heading('Lineas de Producto', level=1)
        linea_table = doc.add_table(rows=len(asesor_data['top_lineas']) + 1, cols=3, style='Table Grid')
        linea_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, h in enumerate(['Linea', 'Cant. Pedida', 'Registros']):
            linea_table.rows[0].cells[idx].text = h
        _style_header_row(linea_table.rows[0], 3, font_size=9)
        for i, linea in enumerate(asesor_data['top_lineas']):
            row = linea_table.rows[i + 1]
            row.cells[0].text = linea['linea']
            row.cells[1].text = f"{linea['cant_pedida']:,.0f}"
            row.cells[2].text = str(linea['registros'])
        _style_alt_rows(linea_table, start_row=1, font_size=9)

    if asesor_data.get('unidades_negocio'):
        doc.add_heading('Unidades de Negocio', level=1)
        co_table = doc.add_table(rows=len(asesor_data['unidades_negocio']) + 1, cols=4, style='Table Grid')
        co_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, h in enumerate(['Unidad', 'Cant. Pedida', 'Valor', 'Registros']):
            co_table.rows[0].cells[idx].text = h
        _style_header_row(co_table.rows[0], 4, font_size=9)
        for i, (co_name, co_data) in enumerate(asesor_data['unidades_negocio'].items()):
            row = co_table.rows[i + 1]
            row.cells[0].text = co_name
            row.cells[1].text = f"{co_data['cant_pedida']:,.0f}"
            row.cells[2].text = f"${co_data['valor']:,.0f}"
            row.cells[3].text = str(co_data['registros'])
        _style_alt_rows(co_table, start_row=1, font_size=9)

        chart_buf = _chart_unidades_negocio(asesor_data)
        if chart_buf:
            doc.add_paragraph()
            doc.add_picture(chart_buf, width=Inches(5.8))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if asesor_data.get('top_estados_produccion'):
        doc.add_heading('Estados de Produccion', level=1)
        est_table = doc.add_table(rows=len(asesor_data['top_estados_produccion']) + 1, cols=2, style='Table Grid')
        est_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, h in enumerate(['Estado', 'Cantidad']):
            est_table.rows[0].cells[idx].text = h
        _style_header_row(est_table.rows[0], 2, font_size=9)
        for i, (estado, cant) in enumerate(asesor_data['top_estados_produccion'].items()):
            row = est_table.rows[i + 1]
            row.cells[0].text = estado
            row.cells[1].text = str(cant)
        _style_alt_rows(est_table, start_row=1, font_size=9)

    if asesor_data.get('top_proyectos'):
        doc.add_heading('Proyectos', level=1)

        chart_buf = _chart_top_proyectos(asesor_data)
        if chart_buf:
            doc.add_picture(chart_buf, width=Inches(5.8))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        docs_por_proy = asesor_data.get('documentos_por_proyecto', {})
        for p in asesor_data['top_proyectos']:
            doc.add_heading(p['proyecto'], level=2)

            summary_table = doc.add_table(rows=2, cols=5, style='Table Grid')
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, h in enumerate(['Pedida', 'Pendiente', 'Comprometida', 'Valor Pend.', 'V.Comprometido']):
                summary_table.rows[0].cells[idx].text = h
            _style_header_row(summary_table.rows[0], 5, font_size=9)
            for idx, v in enumerate([f"{p['cant_pedida']:,.0f}", f"{p['cant_pendiente']:,.0f}",
                                     f"{p['cant_comprometida']:,.0f}", f"${p['valor_pendiente']:,.0f}",
                                     f"${p['v_comprometido']:,.0f}"]):
                summary_table.rows[1].cells[idx].text = v
            _style_alt_rows(summary_table, start_row=1, font_size=9)

            docs = docs_por_proy.get(p['proyecto'], [])
            if not docs:
                continue

            for d in docs:
                doc.add_heading(d['documento'], level=3)

                doc_summary = doc.add_table(rows=2, cols=5, style='Table Grid')
                doc_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, h in enumerate(['Pedida', 'Pendiente', 'Comprometida', 'Valor Pend.', 'V.Comprometido']):
                    doc_summary.rows[0].cells[idx].text = h
                _style_header_row(doc_summary.rows[0], 5, font_size=9)
                for idx, v in enumerate([f"{d['cant_pedida']:,.0f}", f"{d['cant_pendiente']:,.0f}",
                                         f"{d['cant_comprometida']:,.0f}", f"${d['valor_pendiente']:,.0f}",
                                         f"${d['v_comprometido']:,.0f}"]):
                    doc_summary.rows[1].cells[idx].text = v
                _style_alt_rows(doc_summary, start_row=1, font_size=9)

                if d.get('items'):
                    items_table = doc.add_table(rows=len(d['items']) + 1, cols=6, style='Table Grid')
                    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for idx, h in enumerate(['Item', 'Pedida', 'Pendiente', 'Comprom.', 'Val Pend.', 'V.Comprom.']):
                        items_table.rows[0].cells[idx].text = h
                    _style_header_row(items_table.rows[0], 6, font_size=8)
                    for i, it in enumerate(d['items']):
                        row = items_table.rows[i + 1]
                        row.cells[0].text = it['item'][:50]
                        row.cells[1].text = f"{it['cant_pedida']:,.0f}"
                        row.cells[2].text = f"{it['cant_pendiente']:,.0f}"
                        row.cells[3].text = f"{it['cant_comprometida']:,.0f}"
                        row.cells[4].text = f"${it['valor_pendiente']:,.0f}"
                        row.cells[5].text = f"${it['v_comprometido']:,.0f}"
                    _style_alt_rows(items_table, start_row=1, font_size=8)

    if asesor_data.get('desglose_contrato'):
        doc.add_heading('Tipo de Contrato (Instalacion vs Suministro)', level=1)
        cont_table = doc.add_table(rows=len(asesor_data['desglose_contrato']) + 1, cols=5, style='Table Grid')
        cont_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, h in enumerate(['Tipo', 'Pedida', 'Pendiente', 'Comprometida', 'Registros']):
            cont_table.rows[0].cells[idx].text = h
        _style_header_row(cont_table.rows[0], 5, font_size=9)
        for i, (tipo, d) in enumerate(asesor_data['desglose_contrato'].items()):
            row = cont_table.rows[i + 1]
            row.cells[0].text = tipo
            row.cells[1].text = f"{d['pedida']:,.0f}"
            row.cells[2].text = f"{d['pendiente']:,.0f}"
            row.cells[3].text = f"{d['comprometida']:,.0f}"
            row.cells[4].text = str(d['registros'])
        _style_alt_rows(cont_table, start_row=1, font_size=9)

    if informe:
        doc.add_heading('Analisis con Inteligencia Artificial', level=1)
        _add_markdown_to_doc(doc, informe)

    return doc


def _generate_word_bytes(asesor_name: str, canal: str = "") -> bytes:
    result = current_data["result"]
    asesor_data = _get_asesor_data(asesor_name)
    canal_filter = canal or result.get("canal_filter", "")
    cache_key = f"report_{asesor_name.strip().upper()}"
    informe = current_data.get(cache_key, "")

    doc = _build_word_doc(asesor_name, canal_filter, informe, asesor_data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@app.get("/vendors")
async def get_vendors():
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")

    result = current_data["result"]
    vendors = []
    for m in result["asesor_metrics"]:
        name_upper = m["asesor"].strip().upper()
        email = _find_advisor_email(m["asesor"])
        vendors.append({
            "name": m["asesor"],
            "email": email,
            "cant_pedida": m["cant_pedida"],
            "cant_pendiente": m["cant_pendiente"],
            "valor_total": m.get("valor_total", 0),
            "backlog_pct": m["backlog_pct"],
            "documentos_unicos": m["documentos_unicos"],
        })
    return {"vendors": vendors, "canal_filter": result.get("canal_filter", "")}


@app.post("/test-smtp")
async def test_smtp():
    import asyncio
    resend_key = os.getenv("RESEND_API_KEY", "")
    smtp_host = os.getenv("SMTP_HOST", "")

    if resend_key:
        return {"success": True, "message": "Resend API configurada.", "user": resend_key[:8] + "..."}

    if smtp_host:
        try:
            user = await asyncio.to_thread(_test_smtp_connection)
            return {"success": True, "message": "SMTP conectado.", "user": user}
        except Exception as e:
            return {"success": False, "error": str(e)}

    return {"success": False, "error": "No hay RESEND_API_KEY ni SMTP configurado."}


@app.post("/send-emails")
async def send_emails(
    body: list[dict] = Body(...),
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")

    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    canal_global = ""

    results = []

    import asyncio

    for item in body:
        vendor_name = item.get("name", "")
        email = item.get("email", "")
        canal = item.get("canal", "")
        canal_global = canal

        if not vendor_name or not email:
            results.append({"name": vendor_name, "email": email, "success": False, "error": "Sin correo configurado"})
            continue

        cache_key = f"report_{vendor_name.strip().upper()}"
        if cache_key not in current_data and api_key:
            try:
                await asyncio.to_thread(_generate_asesor_report, vendor_name, api_key, canal, provider)
            except Exception:
                pass

        try:
            word_bytes = await asyncio.to_thread(_generate_word_bytes, vendor_name, canal)
            safe_name = vendor_name.replace(' ', '_').replace('.', '').replace('/', '_')
            attachment_name = f"Informe_{safe_name}.docx"

            body_text = f"""Hola {vendor_name},

Adjunto encontras tu informe de ventas.

Canal: {canal or 'Todos'}

---
Este informe fue generado automaticamente por el Sistema de Analisis de Pedidos.
"""
            await asyncio.to_thread(
                _send_email,
                to_email=email,
                subject=f"Informe de Ventas - {vendor_name}",
                body_text=body_text,
                attachment_bytes=word_bytes,
                attachment_name=attachment_name,
            )
            results.append({"name": vendor_name, "email": email, "success": True})
        except Exception as e:
            results.append({"name": vendor_name, "email": email, "success": False, "error": str(e)})

    return {"results": results}


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
